import sys
from docx import Document

src = sys.argv[1]
out = sys.argv[2]
d = Document(src)
with open(out, 'w', encoding='utf-8') as f:
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            f.write(t + "\n")
    # tables
    for ti, tbl in enumerate(d.tables):
        f.write(f"\n--- TABLE {ti+1} ---\n")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            f.write(" | ".join(cells) + "\n")
print("wrote", out)
